"""
Takes .txt file with guest names and  prepared .docx file 
(or paths to them) as command line arguments.
Creates Word document with custom invitations.
"""

import sys
import docx
from typing import List


def invitations(guest_list: List[str], doc_file: str) -> None:
    counter = 0
    for guest_name in guest_list:
        doc = docx.Document(doc_file)
        doc.add_paragraph('It would be pleasure to have the company of', style='normal')
        doc.add_paragraph(guest_name, style='Title')
        doc.add_paragraph('at 11010 Memory Lane on the Evening of', style='normal')
        doc.add_paragraph('April 1st', style='Title')
        doc.add_paragraph("at 7 o'clock", style='normal')
        if counter < len(guest_list):
            doc.add_page_break()
        doc.save('Invitations.docx')
    print('Invitations is ready!')


def main():
    if len(sys.argv) > 2:
        with open(sys.argv[1], 'r') as guest_file:
            # 1st command line argument is guest names file
            guest_list = guest_file.read().splitlines()
        invitations(guest_list, sys.argv[2])
    else:
        print('Arguments not provided.')


if __name__ == '__main__':
    main()
